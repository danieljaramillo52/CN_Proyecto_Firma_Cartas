from docx import Document

# Cargar documento original

#Documentos con reemplazos aceptables.
doc = Document("Controllers/F Carta de Nombramiento - Salario Basico (Sin cambio de salario).docx")

reemplazos = {
  "Nombre Completo del Empleado": "Arthur Kresler Zapata",
  "Número de Documento": "1000.2012.12.1",
  "DD de MM de AAAA": "18 de marzo de 2025",
  "Cargo actual del empleado": "Operario de Logística",
  "Ciudad" : "Medellin",
  "Nombre empleado" : "Abel Ruiz",
  "DD de MM de AA" : "15 de octubre de 2021",
  "$ Salario": "$ 300000",
  "Fecha_dia_actual" : "6 de Agosto 2025",
  "Nombre del cargo al que pasa." : "Operador de Logística" 
}

# Reemplazos dinámicos en cualquier parte del texto.

for p in doc.paragraphs:
    # Unir todos los runs en texto completo
    texto_original = "".join(run.text for run in p.runs)
    texto_modificado = texto_original

    # Aplicar los reemplazos globales
    for clave, nuevo in reemplazos.items():
        if clave in texto_modificado:
            texto_modificado = texto_modificado.replace(clave, nuevo)

    # Si cambió, borramos y reconstruimos el párrafo
    if texto_modificado != texto_original:
        # Tomar las propiedades del primer run como base
        estilo_base = p.runs[0].style if p.runs else None
        fuente_base = p.runs[0].font if p.runs else None

        # Limpiar todos los runs existentes
        for run in p.runs:
            run.clear()

        # Agregar el nuevo texto como un solo run
        nuevo_run = p.add_run(texto_modificado)

        # Restaurar estilo si existía
        if estilo_base:
            nuevo_run.style = estilo_base
        if fuente_base:
            nuevo_run.font.name = fuente_base.name
            nuevo_run.font.size = fuente_base.size
            nuevo_run.font.italic = fuente_base.italic

# Reemplazo accediendo a Tablas para las que aplique. 
#tabla = doc.tables[0]
#
## Datos ordenados por fila (sin encabezado): fila 1 → Básico, fila 2 → Base Variable, fila 3 → Total
#datos = [
#    ("$ 1,000,000", "50%"),
#    ("$ 500,000", "25%"),
#    ("$ 1,500,000", "75%")
#]
#
## Reemplazo directo por posición (sin buscar texto)
#for i, (valor, porcentaje) in enumerate(datos, start=1):  # start=1 para omitir la fila de encabezado
#    tabla.cell(i, 1).text = valor         # Columna "Valores"
#    tabla.cell(i, 2).text = porcentaje    # Columna "% Participación"""
#
#
#
doc.save("doc_actualizado.docx")


