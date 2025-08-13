# Controllers/procesador_formatos.py
import streamlit as st
import re
from docx import Document
from io import BytesIO
from typing import Dict
from Utils.ui_components import ButtonTracker


class ProcesadorFormatosWord:
    """
    Procesa plantillas de Word (python-docx) aplicando reemplazos en texto y, opcionalmente, en tablas.

    - Encapsular la lógica de *copiar* una plantilla cargada en memoria y *aplicar* reemplazos de marcadores en párrafos y tablas.

    Colaboradores / dependencias:
        - `st.session_state["formatos_cartas"]`: diccionario con plantillas `Document` en memoria.
        - `python-docx.Document`: tipo de documento que se manipula.
        - Utiliza métodos internos de clase para separar las operaciones
          (`_replace_in_paragraphs`, `_replace_in_tables`).


    Instnacia ejemplo:
        doc_resultado = ProcesadorFormatosWord().procesar(
            nombre_formato="Carta_Oferta",
            reemplazos={"{{NOMBRE}}": "Ana Pérez", "Salario Básico.": "$2.000.000"},
            contiene_tablas="sí",
        )
    """

    @classmethod
    def _replace_in_paragraphs(cls, doc: Document, reemplazos: Dict[str, str]) -> None:
        """
        Reemplaza cadenas objetivo por sus valores en **todos** los párrafos del documento.

        Responsabilidad:
            - Recorrer párrafos, armar el texto completo, aplicar reemplazos exactos por string,
              y reconstruir el contenido preservando el estilo/fuente del primer run.

        Parámetros:
            doc (Document): Documento de `python-docx` a modificar (in-place).
            reemplazos (Dict[str, str]): Mapa {marcador: valor_nuevo} para sustitución literal.

        Efectos colaterales:
            - Modifica el `doc` recibido en el lugar.
            - El contenido del párrafo se reescribe en un único run; se intenta conservar
              `style` y `font` del primer run original.

        Notas:
            - Si un marcador está fragmentado entre runs, este enfoque lo cubre al trabajar
              sobre el texto completo concatenado; sin embargo, el formato per-run se homogeniza.
        """
        for p in doc.paragraphs:
            texto_original = "".join(run.text for run in p.runs)
            texto_modificado = texto_original
            for clave, nuevo in reemplazos.items():
                if clave in texto_modificado:
                    texto_modificado = texto_modificado.replace(clave, nuevo)

            if texto_modificado != texto_original:
                estilo_base = p.runs[0].style if p.runs else None
                fuente_base = p.runs[0].font if p.runs else None
                
                p.clear()
                
                nuevo_run = p.add_run(texto_modificado)
                if estilo_base:
                    nuevo_run.style = estilo_base
                if fuente_base:
                    nuevo_run.font.name = fuente_base.name
                    nuevo_run.font.size = fuente_base.size
                    nuevo_run.font.italic = fuente_base.italic

    @staticmethod
    def _inject_porcentajes_salario_fijo_variable(reemplazos: Dict[str, str]) -> Dict[str, str]:
        """
        Lee de `reemplazos` las claves **exactas**:
            - "$ Salario Fijo"
            - "$ Salario Variable"
        Calcula sus porcentajes respecto al total y los **inyecta** en:
            - "Salario Fijo %"
            - "Salario Variable %"

        Notas:
            - Se aceptan valores con símbolos/separadores (p. ej. "$ 1.500.000"); se extraen solo dígitos.
            - Si el total es 0, ambos porcentajes se dejan en "0%".
            - Redondeo entero: `round(x*100)` → "NN%".

        Returns:
            Dict[str, str]: El mismo diccionario `reemplazos` con las nuevas claves añadidas/actualizadas.
        """
        to_num = lambda s: float(re.sub(r"[^\d]", "", str(s)) or 0)

        fijo_raw = reemplazos.get("$ Salario Fijo", 0)
        variable_raw = reemplazos.get("$ Salario Variable", 0)

        fijo = to_num(fijo_raw)
        variable = to_num(variable_raw)
        total = fijo + variable

        if total <= 0:
            fijo_pct = "0%"
            variable_pct = "0%"
        else:
            fijo_pct = f"{round((fijo / total) * 100)}%"
            variable_pct = f"{round((variable / total) * 100)}%"

        reemplazos["Salario Fijo %"] = fijo_pct
        reemplazos["Salario Variable %"] = variable_pct
        reemplazos["$ Salario Total"] = f"$ {round(number=total,ndigits=0)}"
        return reemplazos
    
    @staticmethod
    def _inject_horas_semanales(reemplazos: Dict[str, str]):
        """
        Lee de `reemplazos` las clave **exacta**:
            - "REGIONAL FISICA"
        Determina las horas laborales según un único criterio.

        Returns:
            Dict[str, str]: El mismo diccionario `reemplazos` con las nuevas claves añadidas/actualizadas.
        """
        OFICINA_CENRAL = "OFICINA CENTRAL MEDELLIN"
        regional_raw = reemplazos.get("REGIONAL FISICA", 0)
    
        horas = 47 if regional_raw != OFICINA_CENRAL else 45
        
        reemplazos["Num_horas_semanales"] = str(int(horas)) 
        return reemplazos
    
    @classmethod
    def _replace_in_tables(cls, doc: Document, reemplazos: Dict[str, str]) -> None:
        """
        Actualiza valores en la primera tabla con desglose de salario (Básico, Variable, Total).

        Responsabilidad:
            - Tomar valores de `reemplazos` para "Salario Básico." y "Salario Variable.",
              normalizarlos a número, calcular total y porcentajes, y escribirlos en columnas
              1 (valor) y 2 (porcentaje) de las filas 1..3 (encabezado en fila 0).

        Parámetros:
            doc (Document): Documento a modificar (in-place).
            reemplazos (Dict[str, str]): Puede traer importes como texto con símbolos/Separadores.

        Supuestos:
            - La plantilla tiene al menos una tabla y el esquema esperado:
                fila 0: encabezado; filas 1..3: Básico, Variable, Total.
                col 1: valor en COP; col 2: porcentaje.
            - Claves posibles en `reemplazos`: "Salario Básico.", "Salario Variable.".

        """
        if not doc.tables:
            return

        # 1) Tomar valores del dict (acepta "$Salario variable" o "$Salario variable.")
        basico_raw = reemplazos.get("Salario Básico.", 0)
        variable_raw = reemplazos.get("Salario Variable.", 0)

        # 2) Parseo mínimo (quita símbolos y separadores); sin helpers, sin drama
        to_num = lambda s: float(re.sub(r"[^\d]", "", str(s)) or 0)
        
        basico = to_num(basico_raw)
        variable = to_num(variable_raw)
        total = basico + variable

        p_bas = 0 if total == 0 else basico / total
        p_var = 0 if total == 0 else variable / total

        cop = lambda n: "$ " + f"{int(round(n)):,}".replace(",", ".")
        pct = lambda x: f"{round(x*100)}%"

        datos = [
            (cop(basico), pct(p_bas)),   # fila 1: Básico
            (cop(variable), pct(p_var)), # fila 2: Variable
            (cop(total), "100%"),        # fila 3: Total
        ]

        tabla = doc.tables[0]  # si no es la primera, cambia el índice
        for i, (valor, porcentaje) in enumerate(datos, start=1):  # fila 0 = encabezado
            if len(tabla.columns) > 1:
                tabla.cell(i, 1).text = valor
            if len(tabla.columns) > 2:
                tabla.cell(i, 2).text = porcentaje

    

    @classmethod
    def _replace_apl_por_coordenadas(cls, doc: Document, reemplazos: Dict[str, str]) -> None:
        """
        Reemplaza, por coordenadas, los campos específicos del formato
        'F Formalización de APL en Cargo.docx' dentro de la tabla 0.

        Coordenadas (tabla 0):
            - 'Documento_trabajador'            -> (9, 0)
            - 'Documento de identidad líder'    -> (12, 0)

        Notas:
            - Si la celda tiene más texto, se usa .replace(clave, valor).
            - Si falta alguna clave en `reemplazos`, se omite silenciosamente.
        """
        if not doc.tables:
            return

        tabla = doc.tables[0]
        COORDS = {
            "Documento_trabajador": (9, 0),
            "Documento de identidad líder": (12, 0),
        }

        for clave, (fila, col) in COORDS.items():
            if fila < len(tabla.rows) and col < len(tabla.columns) and clave in reemplazos:
                cell = tabla.cell(fila, col)
                # Reemplaza solo la clave, por si la celda tiene prefijos/sufijos (p.ej., "C.C:  ...")
                cell.text = cell.text.replace(clave, str(reemplazos[clave]))



    def procesar(
        self, nombre_formato: str, reemplazos: Dict[str, str], contiene_tablas: str
    ) -> Document:
        """
        Devuelve una **copia** de la plantilla solicitada con reemplazos aplicados en párrafos
        y, opcionalmente, en la primera tabla (desglose salarial).
        """
        FORMATO_F_A_V = "F Otrosi Cambio Salario Fijo a Variable.docx"
        FORMATO_APL = "F Formalización de APL en Cargo.docx"
        
        formatos_mem = st.session_state.get("formatos_cartas", {})
        if nombre_formato not in formatos_mem:
            raise ValueError(f"Formato '{nombre_formato}' no está cargado en memoria.")

        # Crear una copia para no modificar el original
        buffer = BytesIO()
        formatos_mem[nombre_formato].save(buffer)
        buffer.seek(0)
        doc_copia = Document(buffer)

        # Lógica previa de inyección según formato
        if nombre_formato == FORMATO_APL:
            reemplazos = self._inject_horas_semanales(reemplazos)

        if nombre_formato == FORMATO_F_A_V:
            reemplazos = self._inject_porcentajes_salario_fijo_variable(reemplazos)

        # Reemplazos generales en párrafos
        self._replace_in_paragraphs(doc_copia, reemplazos)

        # 👉 Reemplazos por coordenadas SOLO para el formato APL
        if nombre_formato == FORMATO_APL:
            self._replace_apl_por_coordenadas(doc_copia, reemplazos)

        # Lógica de tablas (desglose salarial) solo si aplica
        if contiene_tablas.lower() == "sí":
            self._replace_in_tables(doc_copia, reemplazos)

        return doc_copia

