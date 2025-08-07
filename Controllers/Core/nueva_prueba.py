from docx import Document

# Cargar documento original

#Documentos con reemplazos aceptables.
#doc = Document("F Otrosí cambio de contrato.docx")
doc4 = Document("F Cartas de Nombramiento Salario Variable.docx")
#doc1 = Document("F Otro Si Licencia No Remunerada por Estudio.docx")
#doc2 = Document("F Carta de Nombramiento - Salario Basico (Con cambio de salario).docx")


reemplazos5 = {
    "Nombre Completo del Empleado" : "Arthur Kresller",
    "Nombre del empleado" : "Arthur Kresller",
    "DD de MM de AAAA" : "18 de Marzo de 2025",
    "Ciudad": "Medellin",
    "Número de documento identidad" : "1032.233.243",
    "Fecha de inicio licencia" : "18 de Junio de 2025", 
    "Fecha Final de Licencia": "19 de Junio de 2025",
    "Cargo actual del empleado": "Analista Comercial",
    "Nombre del cargo al que pasa" : "Analista1",
}
tuple1=("Nombre del cargo al que pasa", input("Dígite el nombre del cargo al que pasa el empleado: "))
tuple2=("Fecha inicio", input("Digite la fecha de inicio del empleado en el nuevo cargo: (DD de MM de AAAA) "))

for cada_input in [tuple1, tuple2]:
    reemplazos5[cada_input[0]] = cada_input[1]
  
print(reemplazos5)

# Reemplazos dinámicos en cualquier parte del texto.

for p in doc4.paragraphs:
    # Unir todos los runs en texto completo
    texto_original = "".join(run.text for run in p.runs)
    texto_modificado = texto_original

    # Aplicar los reemplazos globales
    for clave, nuevo in reemplazos5.items():
        if clave in texto_modificado:
            texto_modificado = texto_modificado.replace(clave, nuevo)

    # Si cambió, borramos y reconstruimos el párrafo
    if texto_modificado != texto_original:
        for run in p.runs:
            run.clear()  # limpia sin eliminar el objeto

        # Agregar el nuevo texto como un solo run (o puedes dividir si deseas)
        p.add_run(texto_modificado)


# AReemplazo accediendo a Tablas para las que aplique. 
tabla = doc4.tables[0]

# Datos ordenados por fila (sin encabezado): fila 1 → Básico, fila 2 → Base Variable, fila 3 → Total
datos = [
    ("$ 1,000,000", "50%"),
    ("$ 500,000", "25%"),
    ("$ 1,500,000", "75%")
]

# Reemplazo directo por posición (sin buscar texto)
for i, (valor, porcentaje) in enumerate(datos, start=1):  # start=1 para omitir la fila de encabezado
    tabla.cell(i, 1).text = valor         # Columna "Valores"
    tabla.cell(i, 2).text = porcentaje    # Columna "% Participación"

doc4.save("doc_actualizado.docx")


